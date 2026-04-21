"""
Word文档读取器
从本地Word文件(.docx)读取稿件内容
"""

import os
from typing import List, Dict
from docx import Document


class WordReader:
    """读取本地Word文档"""

    def __init__(self):
        pass

    def read_document(self, file_path: str) -> Dict:
        """
        读取单个Word文档内容

        Args:
            file_path: Word文件路径

        Returns:
            {"title": str, "content": str, "tags": List[str]}
        """
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # 第一个段落作为标题
        title = paragraphs[0] if paragraphs else "无标题"
        content_parts = paragraphs[1:] if len(paragraphs) > 1 else []

        # 提取标签（#开头的内容）
        tags = []
        content_text = []
        for para in content_parts:
            if para.startswith("#"):
                tags.append(para[1:])
            else:
                content_text.append(para)

        return {
            "title": title,
            "content": "\n".join(content_text),
            "tags": tags,
        }

    def read_reference_articles(self, file_path: str) -> List[Dict]:
        """
        从Word文件读取多篇范文

        格式：每篇文章以"标题："开始，"正文："标记内容区域
        "标题："和标题可能在同一行或分开两行

        Args:
            file_path: Word文件路径

        Returns:
            文章列表
        """
        doc = Document(file_path)
        articles = []
        current_article = None
        in_content = False
        pending_title = None  # 暂存的标题（当"标题："和标题分开时）

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if text.startswith("标题：") or text.startswith("标题 :") or text.startswith("标题:"):
                # 保存之前的文章
                if current_article and current_article["content"]:
                    articles.append(current_article)

                # 提取标题（跳过"标题："或"标题 :"或"标题:"前缀）
                if text.startswith("标题："):
                    title_text = text[3:].strip()  # 去掉"标题："
                elif text.startswith("标题 :"):
                    title_text = text[4:].strip()  # 去掉"标题 :"
                elif text.startswith("标题:"):
                    title_text = text[3:].strip()  # 去掉"标题:"
                else:
                    title_text = text
                if title_text:
                    # 标题在同一行
                    current_article = {"title": title_text, "content": [], "tags": []}
                    pending_title = None
                else:
                    # 标题在下一行，先标记pending
                    pending_title = ""
                    current_article = None
                in_content = False
            elif pending_title == "" and current_article is None:
                # 这是pending的标题行
                current_article = {"title": text, "content": [], "tags": []}
                pending_title = None
            elif text in ("正文：", "正文:", "正文 :"):
                if current_article:
                    in_content = True
            elif current_article is not None and in_content:
                # 提取标签（#开头）
                if text.startswith("#"):
                    # 标签可能在一行中有多个，需要拆分
                    tags = text[1:].split("#")
                    for tag in tags:
                        tag = tag.strip()
                        if tag:
                            current_article["tags"].append(tag)
                elif text:
                    current_article["content"].append(text)

        if current_article and current_article["content"]:
            articles.append(current_article)

        return articles